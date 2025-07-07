import os
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import Table, TableStyle
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from num2words import num2words
import win32api
import win32com.client
from datetime import datetime
import ctypes
import ctypes.wintypes

def get_documents_path():
    """Retourne le chemin du dossier Documents de l'utilisateur"""
    # Pour Windows
    if os.name == 'nt':
        CSIDL_PERSONAL = 5  # My Documents
        SHGFP_TYPE_CURRENT = 0  # Get current, not default value
        
        buf = ctypes.create_unicode_buffer(ctypes.wintypes.MAX_PATH)
        ctypes.windll.shell32.SHGetFolderPathW(None, CSIDL_PERSONAL, None, SHGFP_TYPE_CURRENT, buf)
        
        return buf.value
    # Pour macOS et Linux
    else:
        return os.path.join(os.path.expanduser('~'), 'Documents')

# Chemins de base
DOSSIER_EXPORT = os.path.join(get_documents_path(), "Bordereaux Retrait")
DOSSIER_BASE_SCRIPT = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH = os.path.join(DOSSIER_BASE_SCRIPT, "epargne.png")

# Créer les dossiers s'ils n'existent pas
os.makedirs(DOSSIER_EXPORT, exist_ok=True)


def formater_montant(montant):
    """Formate un montant en entier avec séparateur de milliers et devise"""
    return f"{int(round(float(montant))):,}".replace(",", " ") + " FC"

def convertir_en_lettres(montant):
    """Convertit un montant en lettres avec devise complète"""
    montant_entier = int(round(float(montant)))
    mots = num2words(montant_entier, lang='fr').capitalize()
    return mots + " francs congolais"

def exporter_pdf(data):
    # Récupération des données
    nom_complet = data["nom_complet"]
    montant = data["montant_retire"]
    agent = data["agent"]
    ancien_solde = data["ancien_solde"]
    nouveau_solde = data["nouveau_solde"]
    ref = data["ref"]
    numero_client = data["numero_client"]
    numero_carte = data["numero_carte"]
    commission = data.get("commission", 0.0)
    montant_net = float(montant) - float(commission)

    # Formatage des valeurs
    montant_formate = formater_montant(montant)
    ancien_solde_formate = formater_montant(ancien_solde)
    nouveau_solde_formate = formater_montant(nouveau_solde)
    commission_formate = formater_montant(commission)
    montant_net_formate = formater_montant(montant_net)
    montant_lettre = convertir_en_lettres(montant)
    
    date_str = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
    date_signature = datetime.now().strftime('%d/%m/%Y')

    # Création du fichier
    filename = f"{nom_complet.replace(' ', '_')}_retrait_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.pdf"
    path = os.path.join(DOSSIER_EXPORT, filename)

    c = canvas.Canvas(path, pagesize=A4)
    largeur, hauteur = A4
    
    def dessiner_tableau(y_offset, titre):
        y = y_offset
        
        # En-tête avec logo
        if os.path.exists(LOGO_PATH):
            c.drawImage(LOGO_PATH, 2*cm, y - 2*cm, width=3*cm, height=2*cm)

        # Texte d'en-tête
        c.setFont("Helvetica-Bold", 11)
        c.drawCentredString(largeur/2, y - 0.5*cm, "SERVICE CENTRAL D'EPARGNE POUR LA PROMOTION DE L'ENTREPRENEURIAT")
        c.setFont("Helvetica", 9)
        c.drawCentredString(largeur/2, y - 1.2*cm, "$-MONEY / Easy save, Easy get")
        c.drawCentredString(largeur/2, y - 1.8*cm, "Av. Kinsimba n°83bis, Q/Munganga, C/Ngaliema")
        c.drawCentredString(largeur/2, y - 2.4*cm, "Tél: +243 82 058 65 51 / E-mail: save.money.get@gmail.com")

        # Ligne de séparation
        c.setLineWidth(1)
        c.line(2*cm, y - 2.6*cm, largeur - 2*cm, y - 2.6*cm)

        # Titre du bordereau
        c.setFont("Helvetica-Bold", 10)
        c.drawCentredString(largeur/2, y - 3.5*cm, f"BORDEREAU DE RETRAIT ({titre})")

        # Données du tableau
        data_table = [
            ["Nom de l'abonné", nom_complet],
            ["Numéro client", numero_client],
            ["Numéro de la carte", numero_carte],
            ["Montant retiré", montant_formate],
            ["Commission", commission_formate],
            ["Montant net", montant_net_formate],
            ["En lettres", montant_lettre],
            ["Référence retrait", ref],
            ["Ancien solde", ancien_solde_formate],
            ["Nouveau solde", nouveau_solde_formate],
            ["Retrait effectué par", agent],
            ["Date", date_str],
        ]

        # Création du tableau
        table = Table(data_table, colWidths=[7*cm, 8.5*cm])
        style = TableStyle([
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONT', (0, 0), (-1, -1), 'Helvetica', 9),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('ALIGN', (0, 0), (0, -1), 'LEFT'),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('FONT', (0, 3), (-1, 5), 'Helvetica-Bold'),
            ('BACKGROUND', (0, 3), (-1, 5), colors.lightgrey),
        ])
        table.setStyle(style)
        table.wrapOn(c, largeur, hauteur)
        table.drawOn(c, 2*cm, y - 11*cm)

        # Bloc signatures
        signature_y = y - 11.5*cm
        c.setFont("Helvetica", 9)
        c.drawString(12*cm, signature_y + 0.0*cm, f"Fait à Kinshasa, le {date_signature}")
        c.drawString(2*cm, signature_y - 0.7*cm, "Signature de l'abonné : ______________________")
        c.drawString(12*cm, signature_y - 0.7*cm, "Signature de l'agent : ______________________")

        # Ligne de séparation
        c.setLineWidth(0.5)
        c.setStrokeColor(colors.HexColor("#808080"))
        c.line(2*cm, hauteur/2 + 0.3*cm, largeur - 2*cm, hauteur/2 + 0.3*cm)
    
    # Dessin des deux bordures
    dessiner_tableau(hauteur - 2*cm, "ORIGINAL")
    dessiner_tableau(hauteur/2 - 1*cm, "DUPLICATA")

    c.save()
    return path

def exporter_word(data):
    # Récupération des données
    nom_complet = data["nom_complet"]
    montant = data["montant_retire"]
    agent = data["agent"]
    ancien_solde = data["ancien_solde"]
    nouveau_solde = data["nouveau_solde"]
    ref = data["ref"]
    numero_client = data["numero_client"]
    numero_carte = data["numero_carte"]
    commission = data.get("commission", 0.0)
    montant_net = float(montant) - float(commission)

    # Formatage des valeurs
    montant_formate = formater_montant(montant)
    ancien_solde_formate = formater_montant(ancien_solde)
    nouveau_solde_formate = formater_montant(nouveau_solde)
    commission_formate = formater_montant(commission)
    montant_net_formate = formater_montant(montant_net)
    montant_lettre = convertir_en_lettres(montant)
    
    date_str = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
    date_signature = datetime.now().strftime('%d/%m/%Y')

    # Création du fichier
    filename = f"{nom_complet.replace(' ', '_')}_retrait_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.docx"
    path = os.path.join(DOSSIER_EXPORT, filename)

    doc = Document()

    def bordereau(titre):
        if os.path.exists(LOGO_PATH):
            doc.add_picture(LOGO_PATH, width=Inches(1.2))

        # En-tête
        title = doc.add_paragraph("SERVICE CENTRAL D'EPARGNE")
        title.runs[0].bold = True
        doc.add_paragraph("$-MONEY / Easy save, Easy get")
        doc.add_paragraph("Av. Kinsimba n°83bis, Q/Munganga, C/Ngaliema")
        doc.add_paragraph("Tél: +243 82 058 65 51 / E-mail: save.money.get@gmail.com")
        
        # Titre
        subtitle = doc.add_paragraph(f"\nBORDEREAU DE RETRAIT ({titre})")
        subtitle.runs[0].bold = True

        # Données
        lignes = [
            ("Nom de l'abonné", nom_complet),
            ("Numéro client", numero_client),
            ("Numéro de la carte", numero_carte),
            ("Montant retiré", montant_formate),
            ("Commission", commission_formate),
            ("Montant net", montant_net_formate),
            ("En lettres", montant_lettre),
            ("Référence retrait", ref),
            ("Ancien solde", ancien_solde_formate),
            ("Nouveau solde", nouveau_solde_formate),
            ("Retrait effectué par", agent),
            ("Date", date_str),
        ]

        # Création du tableau
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for label, value in lignes:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = value
            
            # Mise en forme des montants
            if label in ["Montant retiré", "Commission", "Montant net"]:
                row[0].paragraphs[0].runs[0].bold = True
                row[1].paragraphs[0].runs[0].bold = True
                row[0].paragraphs[0].runs[0].font.size = Pt(11)
                row[1].paragraphs[0].runs[0].font.size = Pt(11)
        
        # Signatures
        doc.add_paragraph(f"Fait à Kinshasa, le {date_signature}")
        doc.add_paragraph("Signature de l'abonné : ______________________")
        doc.add_paragraph("Signature de l'agent : ______________________")
     
    # Création des deux copies
    bordereau("ORIGINAL")
    doc.add_paragraph("__________________________________________________")
    bordereau("DUPLICATA")

    doc.save(path)
    return path

def imprimer_bordereau(data, commission=0.0):
    """Imprime le bordereau et retourne les chemins des fichiers générés"""
    try:
        # Créer une copie des données avec la commission
        data_with_commission = data.copy()
        data_with_commission["commission"] = commission
        
        # Générer les documents
        pdf_path = exporter_pdf(data_with_commission)
        word_path = exporter_word(data_with_commission)
        
        # Ouvrir les documents pour impression
        try:
            os.startfile(pdf_path)
        except:
            print("Ouverture PDF échouée")
            
        try:
            word = win32com.client.Dispatch("Word.Application")
            docx = word.Documents.Open(word_path)
            docx.PrintOut()
            docx.Close(False)
            word.Quit()
        except Exception as e:
            print("Erreur impression Word:", e)
            
        return pdf_path, word_path

    except Exception as e:
        print(f"Erreur impression bordereau : {e}")
        return None, None