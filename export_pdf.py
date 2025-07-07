import os
from datetime import datetime
from fpdf import FPDF
from num2words import num2words
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
import win32com.client

# --- Chemins pour la compatibilité avec pyinstaller ---
def get_documents_path():
    """Retourne le chemin du dossier Documents de l'utilisateur"""
    # Pour Windows
    if os.name == 'nt':
        import ctypes.wintypes
        CSIDL_PERSONAL = 5  # My Documents
        SHGFP_TYPE_CURRENT = 0  # Get current, not default value
        
        buf = ctypes.create_unicode_buffer(ctypes.wintypes.MAX_PATH)
        ctypes.windll.shell32.SHGetFolderPathW(None, CSIDL_PERSONAL, None, SHGFP_TYPE_CURRENT, buf)
        
        return buf.value
    # Pour macOS et Linux
    else:
        return os.path.join(os.path.expanduser('~'), 'Documents')

# Dossiers de travail
DOSSIER_EXPORT = os.path.join(get_documents_path(), "Bordereaux Epargne")
LOGO_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "epargne.png")

# Créer le dossier s'il n'existe pas
os.makedirs(DOSSIER_EXPORT, exist_ok=True)

# === Fonctions utilitaires ===
def formater_montant(montant):
    """Formate un montant en entier avec séparateur de milliers et devise"""
    return f"{int(round(float(montant))):,}".replace(",", " ") + " FC"

def convertir_en_lettres(montant):
    """Convertit un montant en lettres avec devise complète"""
    montant_entier = int(round(float(montant)))
    mots = num2words(montant_entier, lang='fr').capitalize()
    return mots + " francs congolais"

# === Classe PDF pour créer le bordereau ===
class BordereauPDF(FPDF):
    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=False, margin=10)
        self.add_page()
        self.set_font("Arial", size=9)

    def en_tete(self, y, titre, operation):
        """Affiche l'en-tête du bordereau"""
        # Logo
        if os.path.exists(LOGO_PATH):
            self.image(LOGO_PATH, x=12, y=y, w=15)
        
        # Titre de l'opération
        operation_text = "DÉPÔT" if operation == "depot" else "RETRAIT"
        
        # Coordonnées de l'entreprise
        self.set_xy(30, y)
        self.set_font("Arial", "B", 11)
        self.cell(0, 5, "SERVICE CENTRAL D'EPARGNE", ln=True, align="C")
        
        self.set_font("Arial", "B", 9)
        self.set_xy(30, y + 5)
        self.cell(0, 5, "$-MONEY   |   Easy save, Easy get", ln=True, align="C")
        
        self.set_font("Arial", "", 8)
        self.set_xy(30, y + 10)
        self.cell(0, 5, "Av. Kinsimba n°83 Q/Munganga, C/Ngaliema", ln=True, align="C")
        
        self.set_xy(30, y + 15)
        self.cell(0, 5, "Tél: +243 82 58 65 51   |   E-mail: save.money.get@gmail.com", ln=True, align="C")
        
        # Ligne de séparation
        self.line(10, y + 22, 200, y + 22)
        
        # Titre du bordereau
        self.set_font("Arial", "BU", 10)
        self.set_xy(10, y + 24)
        self.cell(0, 6, f"BORDEREAU DE {operation_text} ({titre})", ln=True, align="C")

    def champ(self, label, valeur):
        """Affiche un champ avec son label et sa valeur"""
        self.set_font("Arial", "B", 8)
        self.set_fill_color(240, 240, 240)
        self.cell(60, 6, label, 1, 0, "L", 1)
        
        self.set_font("Arial", "", 8)
        self.cell(130, 6, str(valeur), 1, 1, "L")

    def generer_bordereau(self, data, y_depart, operation):
        """Génère un bordereau complet"""
        # En-tête
        self.en_tete(y_depart, "ORIGINAL", operation)
        
        # Position de départ pour les champs
        self.set_xy(10, y_depart + 30)
        
        # Formatage des dates
        date_heure = data["date_heure"].split(' ')
        date = date_heure[0] if len(date_heure) > 0 else datetime.now().strftime("%d/%m/%Y")
        heure = date_heure[1][:5] if len(date_heure) > 1 else datetime.now().strftime("%H:%M")
        
        # Champs communs
        self.champ("Nom du client", data["nom_complet"])
        self.champ("Numéro client", data["numero_client"])
        self.champ("Numéro de carte", data["numero_carte"])
        
        # Champs spécifiques à l'opération
        if operation == "depot":
            self.champ("Montant versé", formater_montant(data["montant"]))
        else:  # retrait
            self.champ("Montant retiré", formater_montant(data["montant"]))
        
        # Montant en lettres
        self.champ("En lettres", convertir_en_lettres(data["montant"]))
        
        # Soldes
        self.champ("Ancien solde", formater_montant(data["ancien_solde"]))
        self.champ("Nouveau solde", formater_montant(data["nouveau_solde"]))
        
        # Référence et dates
        self.champ("Référence", data["ref"])
        self.champ("Date", date)
        self.champ("Heure", heure)
        self.champ("Agent en service", data["nom_agent"])
        
        # Signatures
        self.ln(4)
        y = self.get_y()
        
        self.set_font("Arial", "", 8)
        self.cell(90, 6, "Signature de l'abonné: ________________________", ln=0, align="L")
        self.cell(0, 6, f"Fait à Kinshasa, le {date}", ln=1, align="R")
        
        self.cell(90, 6, "", ln=0, align="L")
        self.cell(0, 6, "Signature de l'agent: ________________________", ln=1, align="R")

# === Fonction pour générer le PDF ===
def exporter_bordereau_pdf(data, operation="depot"):
    """Génère un PDF avec deux bordereaux (ORIGINAL et DUPLICATA)"""
    pdf = BordereauPDF()
    
    # Premier bordereau (ORIGINAL) en haut
    pdf.generer_bordereau(data, 10, operation)
    
    # Ligne séparatrice au milieu
    pdf.set_draw_color(150, 150, 150)
    pdf.set_line_width(0.2)
    pdf.line(10, 140, 200, 140)
    
    # Deuxième bordereau (DUPLICATA) en bas
    pdf.generer_bordereau(data, 145, operation)

    # Enregistrer le fichier avec le nom de l'abonné
    horodatage = datetime.now().strftime("%Y%m%d_%H%M%S")
    # Nettoyer le nom pour l'utiliser dans un chemin de fichier
    nom_abonne = data["nom_complet"].replace(" ", "_").replace("/", "-")[:30]  # Limite à 30 caractères
    filename = f"bordereau_{operation}_{nom_abonne}_{horodatage}.pdf"
    chemin_fichier = os.path.join(DOSSIER_EXPORT, filename)
    pdf.output(chemin_fichier)

    return chemin_fichier

# === Fonction pour générer le Word ===
def exporter_bordereau_word(data, operation="depot"):
    """Génère un document Word avec le bordereau"""
    # Formatage des dates
    date_heure = data["date_heure"].split(' ')
    date = date_heure[0] if len(date_heure) > 0 else datetime.now().strftime("%d/%m/%Y")
    heure = date_heure[1][:5] if len(date_heure) > 1 else datetime.now().strftime("%H:%M")
    
    # Création du document
    doc = Document()
    
    # Titre
    operation_text = "DÉPÔT" if operation == "depot" else "RETRAIT"
    title = doc.add_paragraph(f"BORDEREAU DE {operation_text}")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)
    title.alignment = 1  # Centré
    
    # Informations de l'entreprise
    company = doc.add_paragraph("SERVICE CENTRAL D'EPARGNE")
    company.runs[0].bold = True
    company.alignment = 1
    
    doc.add_paragraph("$-MONEY | Easy save, Easy get").alignment = 1
    doc.add_paragraph("Av. Kinsimba n°83 Q/Munganga, C/Ngaliema").alignment = 1
    doc.add_paragraph("Tél: +243 82 58 65 51 | E-mail: save.money.get@gmail.com").alignment = 1
    
    # Tableau des informations
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Fonction pour ajouter une ligne
    def ajouter_ligne(label, valeur):
        row = table.add_row().cells
        row[0].text = label
        row[1].text = valeur
        row[0].paragraphs[0].runs[0].bold = True
    
    # Données communes
    ajouter_ligne("Nom du client", data["nom_complet"])
    ajouter_ligne("Numéro client", data["numero_client"])
    ajouter_ligne("Numéro de carte", data["numero_carte"])
    
    # Données spécifiques
    if operation == "depot":
        ajouter_ligne("Montant versé", formater_montant(data["montant"]))
    else:
        ajouter_ligne("Montant retiré", formater_montant(data["montant"]))
    
    ajouter_ligne("En lettres", convertir_en_lettres(data["montant"]))
    ajouter_ligne("Ancien solde", formater_montant(data["ancien_solde"]))
    ajouter_ligne("Nouveau solde", formater_montant(data["nouveau_solde"]))
    ajouter_ligne("Référence", data["ref"])
    ajouter_ligne("Date", date)
    ajouter_ligne("Heure", heure)
    ajouter_ligne("Agent en service", data["nom_agent"])
    
    # Signatures
    doc.add_paragraph()
    signature_frame = doc.add_paragraph()
    signature_frame.add_run("Signature de l'abonné: ________________________").bold = True
    signature_frame.add_run("\t\t\t\t")
    signature_frame.add_run(f"Fait à Kinshasa, le {date}")
    
    doc.add_paragraph()
    signature_frame2 = doc.add_paragraph()
    signature_frame2.add_run("\t\t\t\tSignature de l'agent: ________________________")
    
    # Séparateur pour le duplicata
    doc.add_paragraph("\n" + "="*100 + "\n")
    
    # Duplicata (copie identique)
    title_dup = doc.add_paragraph(f"DUPLICATA - BORDEREAU DE {operation_text}")
    title_dup.runs[0].bold = True
    title_dup.runs[0].font.size = Pt(14)
    title_dup.alignment = 1
    
    # Copier tout le contenu précédent pour le duplicata
    table_dup = doc.add_table(rows=0, cols=2)
    table_dup.style = 'Table Grid'
    table_dup.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row in table.rows:
        new_row = table_dup.add_row()
        for i, cell in enumerate(row.cells):
            new_row.cells[i].text = cell.text
            new_row.cells[i].paragraphs[0].runs[0].bold = True
    
    # Signatures du duplicata
    doc.add_paragraph()
    signature_frame_dup = doc.add_paragraph()
    signature_frame_dup.add_run("Signature de l'abonné: ________________________").bold = True
    signature_frame_dup.add_run("\t\t\t\t")
    signature_frame_dup.add_run(f"Fait à Kinshasa, le {date}")
    
    doc.add_paragraph()
    signature_frame2_dup = doc.add_paragraph()
    signature_frame2_dup.add_run("\t\t\t\tSignature de l'agent: ________________________")
    
    # Enregistrer le fichier avec le nom de l'abonné
    horodatage = datetime.now().strftime("%Y%m%d_%H%M%S")
    # Nettoyer le nom pour l'utiliser dans un chemin de fichier
    nom_abonne = data["nom_complet"].replace(" ", "_").replace("/", "-")[:30]  # Limite à 30 caractères
    filename = f"bordereau_{operation}_{nom_abonne}_{horodatage}.docx"
    chemin_fichier = os.path.join(DOSSIER_EXPORT, filename)
    doc.save(chemin_fichier)
    
    return chemin_fichier

# === Fonction pour générer les bordereaux SANS impression/ouverture ===
def generer_bordereaux(data, operation="depot"):
    """Génère et enregistre les bordereaux (PDF + Word) sans les ouvrir/imprimer"""
    try:
        # Générer les documents
        pdf_path = exporter_bordereau_pdf(data, operation)
        word_path = exporter_bordereau_word(data, operation)
        return pdf_path, word_path

    except Exception as e:
        print(f"Erreur génération bordereau : {e}")
        return None, None

def exporter_releve_client_pdf(data):
        """Génère un PDF avec le relevé des dépôts d'un client"""
        try:
         # Créer le PDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=9)

            # En-tête avec logo
            if os.path.exists(LOGO_PATH):
                pdf.image(LOGO_PATH, x=12, y=10, w=15)
        
            # Titre
            pdf.set_xy(30, 10)
            pdf.set_font("Arial", "B", 14)
            pdf.cell(0, 8, "RELEVÉ DE DÉPÔTS POUR L'ABONNÉ", ln=True, align="C")
        
            # Informations de l'entreprise
            pdf.set_font("Arial", "B", 9)
            pdf.set_xy(30, 18)
            pdf.cell(0, 5, "SERVICE CENTRAL D'EPARGNE POUR LA PROMOTION DE L'ENTREPRENEURIAT", ln=True, align="C")
        
            pdf.set_font("Arial", "", 8)
            pdf.set_xy(30, 23)
            pdf.cell(0, 5, "$-MONEY   |   Easy save, Easy get", ln=True, align="C")
        
            pdf.set_xy(30, 28)
            pdf.cell(0, 5, "Av. Kinsimba n°83 Q/Munganga, C/Ngaliema", ln=True, align="C")
        
            pdf.set_xy(30, 33)
            pdf.cell(0, 5, "Tél: +243 82 58 65 51   |   E-mail: save.money.get@gmail.com", ln=True, align="C")
        
            # Ligne de séparation
            pdf.line(10, 40, 200, 40)
        
            # Informations du client
            pdf.ln(15)
            pdf.set_font("Arial", "B", 10)
            pdf.cell(40, 6, "Nom complet:", ln=0)
            pdf.set_font("Arial", "", 10)
            pdf.cell(0, 6, data["nom_complet"], ln=1)
        
            pdf.set_font("Arial", "B", 10)
            pdf.cell(40, 6, "Numéro client:", ln=0)
            pdf.set_font("Arial", "", 10)
            pdf.cell(0, 6, data["numero_client"], ln=1)
        
            pdf.set_font("Arial", "B", 10)
            pdf.cell(40, 6, "Date inscription:", ln=0)
            pdf.set_font("Arial", "", 10)
            pdf.cell(0, 6, data["date_inscription"], ln=1)
        
            pdf.set_font("Arial", "B", 10)
            pdf.cell(40, 6, "Durée du compte:", ln=0)
            pdf.set_font("Arial", "", 10)
            pdf.cell(0, 6, data["duree_compte"], ln=1)
        
            # Entête du tableau
            pdf.ln(10)
            pdf.set_font("Arial", "B", 9)
            pdf.set_fill_color(200, 220, 255)  # Couleur d'arrière-plan pour l'en-tête
            pdf.cell(40, 8, "Date", border=1, align="C", fill=True)
            pdf.cell(25, 8, "Heure", border=1, align="C", fill=True)
            pdf.cell(40, 8, "Montant (FC)", border=1, align="C", fill=True)
            pdf.cell(45, 8, "Référence", border=1, align="C", fill=True)
            pdf.cell(40, 8, "Agent", border=1, align="C", fill=True, ln=1)
        
            # Détails des dépôts
            pdf.set_font("Arial", "", 9)
            fill = False
            for depot in data["depots"]:
            # Alternance de couleur pour les lignes
                if fill:
                    pdf.set_fill_color(224, 235, 255)  # Bleu clair
                else:
                    pdf.set_fill_color(255, 255, 255)  # Blanc
                    fill = not fill
            
                pdf.cell(40, 6, depot[0], border=1, align="C", fill=True)
                pdf.cell(25, 6, depot[1], border=1, align="C", fill=True)
                pdf.cell(40, 6, f"{depot[2]:,} FC".replace(",", " "), border=1, align="R", fill=True)
                pdf.cell(45, 6, depot[3], border=1, align="C", fill=True)
                pdf.cell(40, 6, depot[4], border=1, align="C", fill=True, ln=1)
        
            # Total
            pdf.set_font("Arial", "B", 10)
            pdf.set_fill_color(180, 200, 255)  # Couleur différente pour le total
            pdf.cell(105, 8, "TOTAL DES DÉPÔTS:", border=1, align="R", fill=True)
            pdf.cell(40, 8, f"{data['total_depots']:,} FC".replace(",", " "), border=1, align="R", fill=True, ln=1)
        
            # Date de génération
            pdf.ln(10)
            pdf.set_font("Arial", "I", 8)
            pdf.cell(0, 5, f"Relevé généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}", ln=1, align="R")
        
            # Enregistrer le fichier
            nom_abonne = data["nom_complet"].replace(" ", "_").replace("/", "-")[:30]
            horodatage = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"releve_depots_{nom_abonne}_{horodatage}.pdf"
            chemin_fichier = os.path.join(DOSSIER_EXPORT, filename)
            pdf.output(chemin_fichier)
        
            return chemin_fichier
    
        except Exception as e:
            print(f"Erreur génération relevé: {e}")
            return None

# === Alias pour compatibilité ascendante ===
def imprimer_bordereau(data, operation="depot"):
    """Alias pour generer_bordereaux (pour compatibilité) - NE IMPRIME PAS, génère seulement"""
    return generer_bordereaux(data, operation)

# === Fonction de test ===
if __name__ == "__main__":
    # Données de test pour un dépôt
    data_depot = {
        "nom_complet": "John Doe",
        "numero_client": "CLT-001",
        "numero_carte": "CARD-001",
        "montant": 150000,
        "ancien_solde": 200000,
        "nouveau_solde": 350000,
        "ref": "DEP123456",
        "date_heure": datetime.now().strftime("%d/%m/%Y %H:%M"),
        "nom_agent": "Agent Smith"
    }
    
    # Données de test pour un retrait
    data_retrait = {
        "nom_complet": "Jane Smith",
        "numero_client": "CLT-002",
        "numero_carte": "CARD-002",
        "montant": 50000,
        "ancien_solde": 150000,
        "nouveau_solde": 100000,
        "ref": "RET654321",
        "date_heure": datetime.now().strftime("%d/%m/%Y %H:%M"),
        "nom_agent": "Agent Brown"
    }
    
        
    # Tester l'export pour dépôt
    print("Génération bordereau dépôt...")
    pdf_path_depot, word_path_depot = imprimer_bordereau(data_depot, "depot")
    print(f"PDF: {pdf_path_depot}")
    print(f"Word: {word_path_depot}")
    
    # Tester l'export pour retrait
    print("\nGénération bordereau retrait...")
    pdf_path_retrait, word_path_retrait = imprimer_bordereau(data_retrait, "retrait")
    print(f"PDF: {pdf_path_retrait}")
    print(f"Word: {word_path_retrait}")